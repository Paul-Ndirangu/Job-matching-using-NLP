# Cosine Similarity
#
from fastapi import FastAPI, UploadFile, File, Form

from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from typing import Annotated
from docx import Document
import io
from pdfminer.high_level import extract_text 


app = FastAPI()

cv = CountVectorizer()

# Extracting text from PDF files
def extract_text_from_pdf(pdf_path):
    return extract_text(pdf_path)


@app.post("/upload-pdf")
async def upload_pdf(
        job_desc_pdf: UploadFile = File(...),
        res_file_pdf: UploadFile = File(...)
        ):
    jd_pdf = extract_text_from_pdf(job_desc_pdf)
    resume_pdf = extract_text_from_pdf(res_file_pdf)

    jd_pdf = Document(io.BytesIO(), jd_pdf)
    resume_pdf = Document(io.BytesIO(), resume_pdf)
    
    content = [jd_pdf, resume_pdf]
    
    matrix = cv.fit_transform(content)  
    
    similarity_matrix = cosine_similarity(matrix)
    
    match = similarity_matrix[1][0]*100
    
    score = []  
    
    if match > 90:
        score.append(f"Candidate is a match: \nscore: {match}") 
    elif match < 90 and match > 40:
        score.append(f"Candidate is an average match: \nscore: {match}")
    else:
        score.append(f"Candidate is not a match: \nscore: {match}")
    
    return {
            "Job_Description_filename": job_desc_pdf.filename, 
            "Resume_filename": res_file_pdf.filename,
            "Resume_Matches_by": score,
        }

@app.post("/upload-docx")
async def upload_docx(
        job_description_file: UploadFile = File(...),
        resume_file: UploadFile = File(...)
        ):
    
    jd_file = await job_description_file.read()
    file_data = await resume_file.read()  # Read the contents of the uploaded file
    
    
    # Process the file data using python-docx to extract text from the DOCX file
    job_doc = Document(io.BytesIO(jd_file))
    job_description = "\n".join([para.text for para in job_doc.paragraphs])
    
    resume_doc = Document(io.BytesIO(file_data))
    resume = "\n".join([para.text for para in resume_doc.paragraphs])
    
    content = [job_description, resume]
    
    
    matrix = cv.fit_transform(content)
    
    similarity_matrix = cosine_similarity(matrix)
    
    match = similarity_matrix[1][0]*100
    
    score = []
    if match > 90:
        score.append(f"Candidate is a match: \nscore: {match}")
    elif match < 90 and match > 40:
        score.append(f"Candidate is an average match: \nscore: {match}")
    else:
        score.append(f"Candidate is not a match: \nscore: {match}")

    return {
            "Resume_filename": resume_file.filename, 
            "Job_Description": job_description_file.filename,
            "Resume_Matches_by": score,
            "matches": match, 
            "resume": resume
    }
    
   
    